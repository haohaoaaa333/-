#!/usr/bin/env python3
"""Audit a paired Xingce question/answer DOCX without mutating source files."""
from __future__ import annotations

import argparse
import hashlib
import importlib.util
import json
import re
import sys
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def load_parser(script_path: Path):
    spec = importlib.util.spec_from_file_location("xingce_word_parser", script_path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:16]


def docx_inventory(path: Path) -> dict:
    doc = Document(str(path))
    rel_hash = {}
    for rid, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            rel_hash[rid] = {
                "part": str(rel.target_part.partname),
                "hash": sha256(rel.target_part.blob),
                "bytes": len(rel.target_part.blob),
            }

    ordered = []
    image_refs = []
    for index, item in enumerate(doc.element.body.iterchildren()):
        kind = "paragraph" if item.tag.endswith("}p") else "table" if item.tag.endswith("}tbl") else "other"
        texts = [node.text or "" for node in item.xpath(".//w:t")]
        text = "".join(texts).strip()
        rids = [node.get(qn("r:embed")) for node in item.xpath(".//a:blip") if node.get(qn("r:embed"))]
        anchors = len(item.xpath(".//wp:anchor"))
        inlines = len(item.xpath(".//wp:inline"))
        for rid in rids:
            info = dict(rel_hash.get(rid, {"part": "", "hash": "", "bytes": 0}))
            image_refs.append({"block": index, "rid": rid, "anchor": anchors > 0, **info})
        if text or rids:
            ordered.append({
                "index": index,
                "kind": kind,
                "text": text,
                "image_rids": rids,
                "anchors": anchors,
                "inlines": inlines,
            })

    with zipfile.ZipFile(path) as archive:
        media = [entry for entry in archive.infolist() if entry.filename.startswith("word/media/")]

    return {
        "path": str(path),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "body_blocks": len(ordered),
        "relationship_images": len(rel_hash),
        "image_references": len(image_refs),
        "anchored_references": sum(1 for item in image_refs if item["anchor"]),
        "media_files": len(media),
        "media_extensions": dict(Counter(Path(item.filename).suffix.lower() for item in media)),
        "ordered_blocks": ordered,
        "image_refs": image_refs,
    }


def material_signature(question: dict) -> str:
    payload = json.dumps({
        "text": question.get("material", ""),
        "images": question.get("material_images", []),
    }, ensure_ascii=False, sort_keys=True).encode("utf-8")
    return sha256(payload)


def summarize_records(records: list[dict]) -> dict:
    groups = defaultdict(list)
    for question in records:
        if question.get("module_id") == "mod_data":
            groups[material_signature(question)].append(question.get("question_number"))

    def count_nonempty(field: str) -> int:
        return sum(1 for q in records if q.get(field))

    return {
        "questions": len(records),
        "number_min": min((q.get("question_number", 0) for q in records), default=0),
        "number_max": max((q.get("question_number", 0) for q in records), default=0),
        "module_counts": dict(Counter(q.get("module_id") for q in records)),
        "with_material_text": count_nonempty("material"),
        "with_material_images": count_nonempty("material_images"),
        "with_stem_images": count_nonempty("stem_images"),
        "with_option_images": sum(1 for q in records if any(q.get("option_images") or [])),
        "with_explanation_images": count_nonempty("explanation_images"),
        "material_groups": [
            {"signature": signature, "questions": numbers, "size": len(numbers)}
            for signature, numbers in groups.items()
        ],
        "image_counts": {
            "material": sum(len(q.get("material_images", [])) for q in records),
            "stem": sum(len(q.get("stem_images", [])) for q in records),
            "options": sum(sum(len(group) for group in q.get("option_images", [])) for q in records),
            "explanation": sum(len(q.get("explanation_images", [])) for q in records),
        },
    }


def question_marker_diagnostics(inventory: dict) -> dict:
    text = "\n".join(item["text"] for item in inventory["ordered_blocks"] if item["text"])
    sections = [
        ("common", "一、常识判断", "二、言语理解", range(1, 21)),
        ("language", "二、言语理解", "三、数量关系", range(21, 61)),
        ("quantity", "三、数量关系", "四、判断推理", range(61, 71)),
        ("logic", "四、判断推理", "五、资料分析", range(71, 111)),
        ("data", "五、资料分析", "", range(111, 131)),
    ]
    result = {}
    for name, start_marker, end_marker, numbers in sections:
        start = text.find(start_marker)
        end = text.find(end_marker, start + len(start_marker)) if end_marker and start >= 0 else len(text)
        section = text[start:end] if start >= 0 and end >= 0 else ""
        rows = []
        for number in numbers:
            candidates = []
            for match in re.finditer(rf"(?<!\d){number}(?!\d)", section):
                after = section[match.end():match.end() + 24]
                before = section[max(0, match.start() - 20):match.start()]
                next_char = after.lstrip()[:1]
                if next_char and ("\u4e00" <= next_char <= "\u9fff" or next_char in "“《（("):
                    candidates.append({"position": match.start(), "context": (before + f"[{number}]" + after).replace("\n", " ")})
            rows.append({"number": number, "candidate_count": len(candidates), "candidates": candidates[:4]})
        result[name] = {"length": len(section), "markers": rows}
    return result


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--questions", required=True)
    parser.add_argument("--answers", required=True)
    parser.add_argument("--output-dir", required=True)
    args = parser.parse_args()

    question_path = Path(args.questions).resolve()
    answer_path = Path(args.answers).resolve()
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    assets_dir = output_dir / "assets"
    assets_dir.mkdir(exist_ok=True)

    parser_module = load_parser(Path(__file__).with_name("parse_word_bank_with_answers.py"))
    records = parser_module.parse_question_docx(question_path, assets_dir, "/audit-assets")
    answers = parser_module.parse_answer_docx(
        answer_path,
        assets_dir,
        "/audit-assets",
        parser_module.paper_id_for(question_path) + "_answer",
    )
    matched = parser_module.merge_answers(records, answers)

    question_inventory = docx_inventory(question_path)
    answer_inventory = docx_inventory(answer_path)
    report = {
        "question_docx": question_inventory,
        "answer_docx": answer_inventory,
        "parsed": summarize_records(records),
        "answers_found": len(answers),
        "answers_matched": matched,
        "missing_question_numbers": sorted(set(range(1, 131)) - {q.get("question_number") for q in records}),
        "missing_answer_numbers": sorted({q.get("question_number") for q in records} - set(answers)),
        "question_marker_diagnostics": question_marker_diagnostics(question_inventory),
    }

    (output_dir / "records.json").write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "audit.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    summary = {key: value for key, value in report.items() if key not in {"question_docx", "answer_docx", "question_marker_diagnostics"}}
    summary["question_inventory"] = {key: value for key, value in report["question_docx"].items() if key not in {"ordered_blocks", "image_refs"}}
    summary["answer_inventory"] = {key: value for key, value in report["answer_docx"].items() if key not in {"ordered_blocks", "image_refs"}}
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
